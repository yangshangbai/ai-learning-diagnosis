"""题库 Word(docx) 导入/模板导出 + 图片/PDF 云端 OCR 导入接口。

端点：
  POST /api/v1/questions/export-template  前端传入真实题库案例 → 生成含格式说明+带图片的模板 .docx
  POST /api/v1/questions/import-docx      上传 docx → 解析为系统题目 JSON
  POST /api/v1/questions/import-ocr       上传 pdf/png/jpg/jpeg → 云端视觉模型 OCR 识别为题目 JSON

依赖项目根目录的 question_import_export.py（python-docx）。
OCR 依赖：Pillow（图片压缩）+ PyMuPDF(fitz) 或 pypdfium2（PDF 逐页转 PNG，二者装其一即可）。
"""
import base64
import hashlib
import html
import io
import json
import os
import re
import sys
import tempfile
import urllib.error
import urllib.request
import zipfile

from fastapi import APIRouter, Depends, UploadFile, File, Body, Form
from fastapi.responses import Response, JSONResponse
from sqlalchemy.orm import Session

from ..core.app_settings import get_ai_config
from ..core.config import settings
from ..core.db import SessionLocal
from ..core.errors import ValidationError
from ..core.logging import logger
from ..core.security import Principal, require_auth, require_permission
from ..schemas.question import QUES_TYPES

router = APIRouter(prefix="/api/v1/questions", tags=["question-import-export"])


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# 项目根目录（向上三级：routers -> app -> backend -> 项目根）
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".."))
if ROOT not in sys.path:
    sys.path.insert(0, ROOT)

# 前端 demo 可访问的图片目录：生产 nginx 站点根在 frontend/demo，本地开发在 demo/。
# 图片必须落在浏览器能访问的路径下，否则 <img src="images/xxx.png"> 404 导致"有图不显示"。
DEMO_IMG_DIR = None
for _cand in (os.path.join(ROOT, "frontend", "demo", "images"),
              os.path.join(ROOT, "demo", "images")):
    try:
        os.makedirs(_cand, exist_ok=True)
        if DEMO_IMG_DIR is None:
            DEMO_IMG_DIR = _cand
    except Exception:
        continue
if DEMO_IMG_DIR is None:
    DEMO_IMG_DIR = os.path.join(ROOT, "demo", "images")


def _qe():
    import question_import_export as qe
    return qe


_DEFAULT_INSTRUCTIONS = [
    "【格式说明】本模板展示题库导入格式：题目按 1. 2. 3. 编号；每题含题干（可含图片/公式）、【答案】、【解析】。",
    "公式/图形会以图片形式嵌入 Word；按此格式整理后，可在「题库管理 → 导入管理」上传 .docx 导入系统。",
    "以下是真实题库案例（含图片）：",
]


@router.post("/export-template")
def export_template(
    body: dict = Body(...),
    _: Principal = Depends(require_auth),
):
    """生成"人手可操作"的表格导入模板（表头 + 示例行 + 空行），非 IT 人员也能直接填写。"""
    qe = _qe()
    title = (body.get("title") or "题库导入模板").strip() or "题库导入模板"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        qe.export_table_template_docx(tmp_path, title=title)
        with open(tmp_path, "rb") as f:
            content = f.read()
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="question_import_template.docx"'},
    )


@router.post("/import-docx")
async def import_docx(
    file: UploadFile = File(...),
    subject: str = "",
    grade: str = "",
    difficulty: int = 3,
    _: Principal = Depends(require_permission("sync", "add")),
):
    """上传 docx，解析为题目列表返回（docx 内嵌图片直接提取为文件，并以 <img> 插入题干内容）。"""
    if not (file.filename or "").lower().endswith(".docx"):
        return JSONResponse({"code": 400, "message": "仅支持 .docx 文件", "data": None}, 400)
    data = await file.read()
    qe = _qe()
    src_hash = hashlib.md5(data).hexdigest()[:6].upper()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    result = None
    try:
        qs = qe.parse_docx(tmp_path, img_dir=DEMO_IMG_DIR,
                           subject=subject, grade=grade, difficulty=difficulty,
                           source="docx", category=subject, src_prefix="/images/",
                           src_hash=src_hash)
        # 图片已在解析时以 <img src="images/xxx.png"> 直接插入题干（docx 内嵌图 → 文件 → 题干内联）
        img_total = sum(1 for q in qs if "<img" in (q.get("stem") or ""))
        result = {"code": 0, "message": "ok",
                  "data": {"questions": qs, "count": len(qs), "img_questions": img_total}}
    except Exception as e:
        # 伪装扩展名/损坏文件：python-docx 解析失败应回 4xx 而非 500（BUG-L013）
        result = JSONResponse(
            {"code": 400, "message": "文件无法解析为 docx 题库，请检查文件是否完整（%s）" % str(e)[:80], "data": None},
            400,
        )
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass
    return result


# ================================================================
# 云端 OCR 导入（PDF/PNG/JPG → 视觉模型识别 → 题目 JSON）
# ================================================================

# OpenAI 兼容视觉模型 provider 端点表（与 demo「AI模型配置」预设一致）
_VISION_PROVIDERS = {
    "zhipu": "https://open.bigmodel.cn/api/paas/v4/chat/completions",       # 智谱 GLM-4V
    "qwen": "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions",  # 通义 Qwen-VL
    "openai": "https://api.openai.com/v1/chat/completions",                 # OpenAI GPT-4o
    "moonshot": "https://api.moonshot.cn/v1/chat/completions",              # 月之暗面 Kimi-Vision
}
_OCR_EXT = (".pdf", ".png", ".jpg", ".jpeg")
_OCR_MAX_SIDE = 2048      # 图片最长边像素上限


def _build_ocr_prompt(subject: str, grade: str) -> str:
    """构造 OCR 识别 prompt：要求模型严格输出 JSON 数组；含图题标注图片位置，不解析图片内容。"""
    return (
        "你是题库录入专家。请识别图片中的每一道题目（含题干、选项、答案、解析），"
        "并输出为 JSON 数组。\n"
        f"学科：{subject or '（未知）'}，年级：{grade or '（未知）'}\n\n"
        "每道题格式：\n"
        '{"stem":"题干文字","ques_type":"single_choice或multi_choice或fill_blank或true_false或essay",'
        '"options":["A. 选项内容","B. 选项内容","C. 选项内容","D. 选项内容"],'
        '"answer":"答案","analysis":"解析","difficulty":3,'
        '"img_box":null}\n\n'
        "规则：\n"
        "1. 选择题（单选/多选）必须完整识别 A/B/C/D 四个选项与题干；多选题 answer 写选项字母组合（如 AC）。\n"
        "2. 判断题 answer 写 正确/错误 或 对/错；填空题 answer 写填空内容；解答题 answer 写参考答案要点。\n"
        "3. analysis 为解析，识别不到填空字符串。\n"
        "4. difficulty 为 1~5 整数（1 最简单、5 最难），识别不到用 3。\n"
        "5. 严格只输出一个 JSON 数组（数组元素是题目对象），不要输出任何其他文字、解释或 markdown 代码块标记。\n"
        "6. 重要——图片处理：题目里若含有图形/示意图/数轴/坐标系等图片（题干中出现\"如图\"或明显有插图），"
        "【不要解析图片内容成文字，也不要描述图片】；只需估计该图片在这张页面图片中的位置范围，"
        "输出 img_box 为 [x1, y1, x2, y2]（整数，左上角与右下角像素坐标，范围 0~图片宽/高）。"
        "没有图片的题 img_box 输出 null。系统会把该区域原样裁剪作为题目图片插入。"
    )


def _pdf_to_images(data: bytes, max_side: int = _OCR_MAX_SIDE):
    """PDF 逐页渲染为 PNG 图片（优先 PyMuPDF，其次 pypdfium2）。

    返回 list[bytes]，每项为一页 PNG；两者都未安装时抛出清晰错误。
    """
    try:
        import pymupdf as fitz  # PyMuPDF 新版命名（fitz 将废弃）
    except ImportError:
        try:
            import fitz  # PyMuPDF 旧版命名
        except ImportError:
            fitz = None
    if fitz is not None:
        try:
            doc = fitz.open(stream=data, filetype="pdf")
        except Exception as e:
            raise ValidationError("PDF 解析失败：" + str(e))
        try:
            pages = []
            for page in doc:
                # 2x 缩放保证小字可读；超长边再降采样
                zoom = 2.0
                pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
                if pix.width > max_side or pix.height > max_side:
                    s = max_side / max(pix.width, pix.height)
                    pix = page.get_pixmap(matrix=fitz.Matrix(zoom * s, zoom * s))
                pages.append(pix.tobytes("png"))
            return pages
        finally:
            doc.close()
    try:
        import pypdfium2 as pdfium
    except ImportError:
        raise ValidationError(
            "PDF 识别依赖未安装：请安装 PyMuPDF 或 pypdfium2（pip install PyMuPDF，生产同）"
        )
    try:
        pdf = pdfium.PdfDocument(data)
    except Exception as e:
        raise ValidationError("PDF 解析失败：" + str(e))
    try:
        pages = []
        for page in pdf:
            # 渲染到最长边 ~max_side 且不低于 2x（保证清晰度）
            scale = min(max_side / max(page.get_size()), 2.0)
            bitmap = page.render(scale=scale)
            pil_img = bitmap.to_pil().convert("RGB")
            buf = io.BytesIO()
            pil_img.save(buf, format="PNG")
            pages.append(buf.getvalue())
        return pages
    finally:
        pdf.close()


def _normalize_image(data: bytes, max_side: int = _OCR_MAX_SIDE) -> bytes:
    """图片压缩：最长边压到 max_side、转 JPEG，控制上传体量。"""
    try:
        from PIL import Image
    except ImportError:
        raise ValidationError("图片处理依赖未安装：pip install Pillow")
    try:
        img = Image.open(io.BytesIO(data))
        img.load()
    except Exception as e:
        raise ValidationError("图片解析失败：" + str(e))
    img = img.convert("RGB")  # 去掉 alpha，兼容所有视觉模型
    w, h = img.size
    if max(w, h) > max_side:
        s = max_side / max(w, h)
        img = img.resize((int(w * s), int(h * s)), Image.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=88)
    return buf.getvalue()


def _call_vision_model(url: str, api_key: str, model: str, image_bytes: bytes, mime: str, prompt: str) -> str:
    """调用 OpenAI 兼容视觉模型（/chat/completions，messages 内 image_url），返回文本。"""
    img_url = "data:%s;base64,%s" % (mime, base64.b64encode(image_bytes).decode("ascii"))
    payload = {
        "model": model,
        "messages": [{
            "role": "user",
            "content": [
                {"type": "image_url", "image_url": {"url": img_url}},
                {"type": "text", "text": prompt},
            ],
        }],
        "temperature": 0.2,
    }
    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Authorization": "Bearer " + api_key, "Content-Type": "application/json"},
    )
    try:
        with urllib.request.urlopen(req, timeout=180) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        detail = e.read().decode("utf-8", "ignore")[:300] if e.fp else str(e)
        raise ValidationError("视觉模型请求失败（HTTP %s）：%s" % (e.code, detail))
    except urllib.error.URLError as e:
        raise ValidationError("视觉模型请求失败：" + str(e.reason))
    except TimeoutError:
        raise ValidationError("视觉模型请求超时，请重试或检查网络")
    if not data.get("choices"):
        raise ValidationError("模型未返回有效结果：" + json.dumps(data, ensure_ascii=False)[:300])
    return data["choices"][0]["message"]["content"] or ""


def _extract_balanced(text: str, start: int, open_ch: str, close_ch: str):
    """从 start 位置开始做括号平衡匹配，返回完整 JSON 片段或 None。"""
    depth = 0
    in_str = False
    esc = False
    for i in range(start, len(text)):
        ch = text[i]
        if in_str:
            if esc:
                esc = False
            elif ch == "\\":
                esc = True
            elif ch == '"':
                in_str = False
            continue
        if ch == '"':
            in_str = True
        elif ch == open_ch:
            depth += 1
        elif ch == close_ch:
            depth -= 1
            if depth == 0:
                return text[start:i + 1]
    return None


def _parse_json_object(text: str):
    """从模型输出中提取单个 JSON 对象（容忍 ```json``` 包裹、前后杂文）。"""
    if not text:
        return {}
    t = text.strip()
    m = re.search(r"```(?:json)?\s*(\{[\s\S]*?\})\s*```", t)
    if m:
        t = m.group(1)
    try:
        v = json.loads(t)
        return v if isinstance(v, dict) else {}
    except Exception:
        pass
    start = t.find("{")
    if start >= 0:
        seg = _extract_balanced(t, start, "{", "}")
        if seg:
            try:
                v = json.loads(seg)
                return v if isinstance(v, dict) else {}
            except Exception:
                pass
    return {}


def _parse_json_array(text: str):
    """从模型输出中提取 JSON 数组（容错：markdown 包裹 / 直接数组 / 包在对象里）。"""
    if not text:
        return []
    t = text.strip()
    m = re.search(r"```(?:json)?\s*(\[[\s\S]*?\])\s*```", t)
    if m:
        t = m.group(1)
    # 1. 直接是数组
    try:
        v = json.loads(t)
        if isinstance(v, list):
            return v
    except Exception:
        pass
    # 2. 对象里带 questions 等数组字段
    obj = _parse_json_object(t)
    if obj:
        for k in ("questions", "items", "data", "list"):
            v = obj.get(k)
            if isinstance(v, list):
                return v
        for v in obj.values():
            if isinstance(v, list):
                return v
    # 3. 括号平衡匹配第一个完整数组
    start = t.find("[")
    if start >= 0:
        seg = _extract_balanced(t, start, "[", "]")
        if seg:
            try:
                v = json.loads(seg)
                if isinstance(v, list):
                    return v
            except Exception:
                pass
    return []


_TYPE_ALIAS = {
    "single_choice": "single_choice", "single": "single_choice", "单选": "single_choice", "选择题": "single_choice",
    "multi_choice": "multi_choice", "multiple_choice": "multi_choice", "multi": "multi_choice", "多选": "multi_choice", "不定项": "multi_choice",
    "fill_blank": "fill_blank", "fill": "fill_blank", "填空": "fill_blank", "填空题": "fill_blank",
    "true_false": "true_false", "judge": "true_false", "判断": "true_false", "判断题": "true_false",
    "essay": "essay", "解答": "essay", "解答题": "essay", "简答": "essay", "简答题": "essay", "问答": "essay",
}


def _normalize_ocr_question(q, idx: int, subject: str, grade: str, default_diff: int) -> dict:
    """把模型返回的单题 dict 归一化为系统题目 JSON（与 parse_docx 同构，source=ocr）。"""
    if not isinstance(q, dict):
        return None
    stem = str(q.get("stem") or q.get("题干") or "").strip()
    if not stem:
        return None
    qtype_raw = str(q.get("ques_type") or q.get("type") or q.get("题型") or "essay").strip()
    qtype = _TYPE_ALIAS.get(qtype_raw)
    if qtype not in QUES_TYPES:
        qtype = qtype_raw if qtype_raw in QUES_TYPES else "essay"
    # 选择题：优先 options 数组，其次从题干/答案反推
    options = []
    opts_raw = q.get("options") or []
    if qtype in ("single_choice", "multi_choice"):
        for o in opts_raw:
            if isinstance(o, dict):
                label = str(o.get("label") or "").strip()
                content = str(o.get("content") or o.get("text") or "").strip()
                if content:
                    options.append((label + ". " + content).strip() if label else content)
            else:
                s = str(o).strip()
                if s:
                    options.append(s)
        if not options:
            options = re.findall(r"(?:^|[\s\n])([A-D])[\.、．]\s*(.+?)(?=(?:\s|$)(?:[A-D])[\.、．]|$)", stem, re.S)
            options = [("%s. %s" % (a, b)).strip() for a, b in options]
        # 去重保序
        seen = set()
        options = [o for o in options if not (o in seen or seen.add(o))]
    # answer / analysis 归一化
    def _to_text(v):
        if isinstance(v, list):
            return "\n".join(_to_text(x) for x in v if x not in (None, ""))
        if isinstance(v, dict):
            return json.dumps(v, ensure_ascii=False)
        return str(v or "").strip()

    answer = _to_text(q.get("answer") or q.get("答案"))
    analysis = _to_text(q.get("analysis") or q.get("解析"))
    try:
        diff = int(q.get("difficulty") or default_diff)
    except (TypeError, ValueError):
        diff = default_diff
    if diff < 1 or diff > 5:
        diff = 3
    qe = _qe()
    return {
        "id": "OCR%03d" % idx,
        "code": qe._gen_code(subject, grade, idx),
        "subject": subject,
        "grade": grade,
        "type": qtype,
        "difficulty": diff,
        "score": 0,
        "stem": "<p>" + html.escape(stem).replace("\n", "</p><p>") + "</p>",
        "options": options or None,
        "answer": ("<p>" + html.escape(answer).replace("\n", "</p><p>") + "</p>") if answer else "",
        "analysis": ("<p>" + html.escape(analysis).replace("\n", "</p><p>") + "</p>") if analysis else "",
        "knowledge": [],
        "category": subject,
        "source": "ocr",
        "images": [],
    }


@router.post("/import-ocr")
async def import_ocr(
    file: UploadFile = File(...),
    subject: str = Form(""),
    grade: str = Form(""),
    difficulty: int = Form(3),
    provider: str = Form("zhipu"),
    api_key: str = Form(""),
    model: str = Form("glm-4v"),
    base_url: str = Form(""),
    _: Principal = Depends(require_permission("sync", "add")),
    db: Session = Depends(get_db),
):
    """上传 PDF/PNG/JPG → 云端视觉模型 OCR 识别为题目列表返回（source=ocr）。

    - PDF：逐页渲染为 PNG 后逐页识别，合并结果；
    - PNG/JPG：直接压缩后交给视觉模型；
    - 模型凭据优先取数据库设置（系统设置→AI模型配置·视觉模型），env 兜底；
      前端随请求传参仍兼容。
    """
    filename = (file.filename or "").lower()
    ext = os.path.splitext(filename)[1].lower()
    if ext not in _OCR_EXT:
        return JSONResponse(
            {"code": 400, "message": "仅支持 .pdf/.png/.jpg/.jpeg 文件", "data": None}, 400
        )
    data = await file.read()
    if not data:
        return JSONResponse({"code": 400, "message": "上传文件为空", "data": None}, 400)
    aicfg = get_ai_config(db)
    vcfg = aicfg.get("vision") or {}
    if "****" in (api_key or ""):   # 前端掩码回显值误传：视为未传
        api_key = ""
    if not (api_key or "").strip():
        if vcfg.get("api_key"):
            api_key = vcfg["api_key"]   # 数据库设置（env 已在 get_ai_config 内回退）
        else:
            raise ValidationError("未配置视觉模型 API Key（系统设置→AI模型配置，或服务端 AI_ZHIPU_API_KEY）")
    model = (vcfg.get("model") or "").strip() or model

    url = (base_url or "").strip() or _VISION_PROVIDERS.get((provider or "").lower())
    if not url:
        raise ValidationError(
            "不支持的视觉模型 provider：%s（可传 base_url 覆盖）" % provider
        )

    # 1. 文件 → 图片列表
    if ext == ".pdf":
        images = [("image/png", b) for b in _pdf_to_images(data)]
    else:
        mime = "image/png" if ext == ".png" else "image/jpeg"
        images = [(mime, _normalize_image(data))]

    prompt = _build_ocr_prompt(subject, grade)

    # 2. 逐图调用视觉模型识别，合并题目
    raw_list = []
    for i, (mime, img_bytes) in enumerate(images, 1):
        logger.info("ocr_recognize_page", extra={"page": i, "total": len(images)})
        raw = _call_vision_model(url, api_key, model, img_bytes, mime, prompt)
        raw_list.append(raw)

    questions = []
    counter = 0
    for i, raw in enumerate(raw_list):
        mime, img_bytes = images[i] if i < len(images) else ("image/png", b"")
        for q in _parse_json_array(raw):
            counter += 1
            nq = _normalize_ocr_question(q, counter, subject, grade, difficulty)
            if not nq:
                continue
            # 含图题：按模型标注的 img_box 从原页裁剪图片，保存并插入题干（图片原样，不解析内容）
            box = q.get("img_box") or q.get("image_box") or None
            if box and img_bytes:
                try:
                    box = [int(x) for x in box]
                    if len(box) == 4 and box[0] < box[2] and box[1] < box[3]:
                        from PIL import Image as _PIL
                        import io as _io
                        pil = _PIL.open(_io.BytesIO(img_bytes))
                        w, h = pil.size
                        # 约束到图片范围内，适当外扩 6px 保证完整
                        x1 = max(0, box[0] - 6); y1 = max(0, box[1] - 6)
                        x2 = min(w, box[2] + 6); y2 = min(h, box[3] + 6)
                        if x2 - x1 > 10 and y2 - y1 > 10:
                            crop = pil.crop((x1, y1, x2, y2))
                            os.makedirs(DEMO_IMG_DIR, exist_ok=True)
                            fname = "ocr_img_%04d.png" % counter
                            crop.save(os.path.join(DEMO_IMG_DIR, fname), "PNG")
                            nq["stem"] = (nq.get("stem") or "") + (
                                '<p><img src="/images/%s" style="max-width:100%%;"></p>' % fname)
                            nq.setdefault("images", []).append(fname)
                except Exception as e:
                    logger.warning("ocr_crop_fail", extra={"err": str(e)})
            questions.append(nq)

    if not questions:
        return JSONResponse(
            {"code": 422, "message": "未能识别出任何题目（请检查图片清晰度 / 模型是否支持视觉）", "data": None}, 422
        )
    logger.info("ocr_import_done", extra={"file": filename, "pages": len(images), "questions": len(questions)})
    return {"code": 0, "message": "ok", "data": {"questions": questions, "count": len(questions)}}


# ================================================================
# 智能导入（import-smart）：自动识别格式 → 自动路由解析 → 智能推断学科/年级
# ================================================================
_SMART_MAX_SIZE = 30 * 1024 * 1024   # 单文件上限 30MB

# 学科 / 年级关键词（文件名 + 文件内容推断共用）
_SUBJECT_KEYWORDS = ["数学", "语文", "英语", "物理", "化学", "生物", "历史", "地理", "政治", "道德与法治", "科学"]
# (关键词, 标准年级名) —— 七年级→初一、八年级→初二、九年级→初三 归一
_GRADE_KEYWORDS = [
    ("一年级", "一年级"), ("二年级", "二年级"), ("三年级", "三年级"),
    ("四年级", "四年级"), ("五年级", "五年级"), ("六年级", "六年级"),
    ("七年级", "初一"), ("初一", "初一"),
    ("八年级", "初二"), ("初二", "初二"),
    ("九年级", "初三"), ("初三", "初三"),
    ("高一", "高一"), ("高二", "高二"), ("高三", "高三"),
]


def _decode_text(data: bytes) -> str:
    """按 utf-8-sig → utf-8 → gb18030 → latin-1 逐级解码文本；失败抛错。"""
    if not data:
        raise ValueError("empty")
    # 二进制特征：出现大量 NUL 字节 → 不是文本
    if data.count(b"\x00") > max(len(data), 1) * 0.02:
        raise ValueError("binary content")
    errors = []
    for enc in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            text = data.decode(enc)
        except (UnicodeDecodeError, LookupError) as e:
            errors.append("%s:%s" % (enc, str(e)))
            continue
        # 可读性校验：控制字符（除 \n \t \r）过多 → 视为二进制
        head = text[:4000]
        bad = sum(1 for ch in head if ord(ch) < 32 and ch not in "\n\t\r")
        if bad > len(head) * 0.05:
            errors.append("%s:too_many_ctrl" % enc)
            continue
        return text
    raise ValueError("无法解码文本内容：" + "; ".join(errors))


def _detect_file_format(data: bytes, filename: str) -> str:
    """根据 magic bytes + zip 内部结构判定真实格式（扩展名只作提示，防伪装扩展名）。

    返回格式 key：docx / xlsx / pdf / png / jpg / text；识别不了抛 ValidationError。
    """
    head = data[:16]
    # OLE 老版 Word(.doc) / Excel(.xls)
    if head[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
        raise ValidationError("检测到旧版 Office(.doc/.xls) 二进制格式，请用 Word/Excel 另存为 .docx/.xlsx 后再导入")
    # RTF
    if head[:5].lower() == b"{\\rtf":
        raise ValidationError("检测到 RTF 格式，请用 Word 另存为 .docx 后再导入")
    # ZIP 容器 → docx / xlsx / 其他
    if data[:4] == b"PK\x03\x04" or data[:4] == b"PK\x05\x06":
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                names = [n.lower() for n in zf.namelist()]
        except Exception:
            raise ValidationError("文件已损坏或不是有效的 Office 文档（压缩结构读取失败），请重新导出后再导入")
        if any(n.startswith("word/") for n in names):
            return "docx"
        if any(n.startswith("xl/") for n in names):
            return "xlsx"
        raise ValidationError("不支持的文件类型：这是一个压缩包，但不是 Word(.docx) 或 Excel(.xlsx) 文档")
    if head[:5] == b"%PDF-":
        return "pdf"
    if head[:8] == b"\x89PNG\r\n\x1a\n":
        return "png"
    if head[:3] == b"\xff\xd8\xff":
        return "jpg"
    # 兜底：尝试按文本解码（txt/csv/md 及伪装扩展名的文本文件）
    try:
        _decode_text(data)
        return "text"
    except ValueError as e:
        raise ValidationError("无法识别的文件类型（%s）。支持的格式：Word(.docx)/PDF/图片(PNG/JPG)/Excel(.xlsx)/文本(.txt/.csv/.md)" % e)


def _smart_infer_meta(filename: str, content: str = "",
                      template_subject: str = "", template_grade: str = "") -> tuple:
    """智能推断学科/年级。优先级：模板列 > 正文首部 > 文件名。返回 (subject, grade)。"""
    subject = (template_subject or "").strip()
    grade = (template_grade or "").strip()

    head = (content or "")[:3000]
    if not subject:
        for kw in _SUBJECT_KEYWORDS:
            if kw in head:
                subject = kw
                break
    if not grade:
        for kw, std in _GRADE_KEYWORDS:
            if kw in head:
                grade = std
                break

    name = filename or ""
    if not subject:
        for kw in _SUBJECT_KEYWORDS:
            if kw in name:
                subject = kw
                break
    if not grade:
        for kw, std in _GRADE_KEYWORDS:
            if kw in name:
                grade = std
                break
    return subject, grade


def _extract_docx_text(path: str) -> str:
    """提取 docx 纯文本（段落 + 表格单元格），供学科/年级内容推断。失败返回空串。"""
    try:
        from docx import Document as _Doc
        doc = _Doc(path)
        parts = [p.text for p in doc.paragraphs]
        for tbl in doc.tables:
            for row in tbl.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    except Exception:
        return ""


def _strip_md(line: str) -> str:
    """剥离常见 Markdown 标记，保留可读文本。"""
    s = line.strip()
    s = re.sub(r"^#{1,6}\s*", "", s)      # 标题 #
    s = re.sub(r"^>\s*", "", s)           # 引用 >
    s = re.sub(r"^\*\s*", "", s)          # 无序列表 *
    s = s.replace("**", "").replace("`", "").replace("~~", "")  # 注意：不能去 __，填空题下划线靠它识别
    return s.strip()


_ANS_RE = re.compile(r"^[【\[]\s*答案\s*[】\]]\s*[:：]?\s*(.*)$")
_ANA_RE = re.compile(r"^[【\[]\s*解析\s*[】\]]\s*[:：]?\s*(.*)$")
_ANS_PLAIN_RE = re.compile(r"^答案\s*[:：]\s*(.*)$")
_ANA_PLAIN_RE = re.compile(r"^解析\s*[:：]\s*(.*)$")
_OPT_RE = re.compile(r"^\s*([A-Ea-e])\s*[\.、．)）]\s*(.*)$")


def _parse_text(text: str, subject: str = "", grade: str = "", difficulty: int = 3,
                source: str = "text") -> list:
    """文本（txt/csv/md）→ 题目列表。先试 csv 列式模板，否则按编号切分。"""
    if not text or not text.strip():
        return []
    # 1) csv 列式模板（表头含 题型+题干）
    col_qs = _parse_text_columns(text, subject, grade, difficulty, source)
    if col_qs is not None:
        return col_qs
    # 2) 编号切分
    questions = []
    cur = None
    num_re = re.compile(r"^\s*(\d{1,3})\s*[\.、．)）]\s*")
    for raw in text.splitlines():
        line = _strip_md(raw)
        if not line:
            continue
        m = num_re.match(line)
        if m:
            if cur and (cur["stem"] or cur["answer"]):
                questions.append(cur)
            cur = {"num": int(m.group(1)), "stem": [], "options": [], "answer": [], "analysis": []}
            rest = line[m.end():].strip()
            if rest:
                cur["stem"].append(rest)
            continue
        if cur is None:
            continue  # 编号前的标题/说明行跳过
        for rex, key in ((_ANS_RE, "answer"), (_ANS_PLAIN_RE, "answer"),
                         (_ANA_RE, "analysis"), (_ANA_PLAIN_RE, "analysis")):
            mm = rex.match(line)
            if mm:
                cur[key].append(mm.group(1).strip())
                break
        else:
            om = _OPT_RE.match(line)
            if om and not cur["answer"] and not cur["analysis"]:
                cur["options"].append("%s. %s" % (om.group(1).upper(), om.group(2).strip()))
            elif not cur["answer"] and not cur["analysis"]:
                cur["stem"].append(line)
            elif cur["analysis"]:
                cur["analysis"].append(line)
            else:
                cur["answer"].append(line)
    if cur and (cur["stem"] or cur["answer"]):
        questions.append(cur)

    qe = _qe()
    result = []
    for i, q in enumerate(questions, 1):
        stem_text = " ".join(q["stem"]).strip()
        if not stem_text:
            continue
        # 选项文本并入判定源（如 A.正确/B.错误 需据此判判断题），但不写入题干
        type_text = stem_text + " " + " ".join(q["options"])
        typ = qe.detect_type(type_text, q["options"]) if hasattr(qe, "detect_type") else "essay"
        answer = " ".join(q["answer"]).strip()
        analysis = " ".join(q["analysis"]).strip()
        result.append({
            "id": "TXT%03d" % i,
            "code": qe._gen_code(subject, grade, i),
            "subject": subject,
            "grade": grade,
            "type": typ,
            "difficulty": difficulty,
            "score": 0,
            "stem": "<p>" + html.escape(stem_text).replace("\n", "</p><p>") + "</p>",
            "options": q["options"] or None,
            "answer": ("<p>" + html.escape(answer).replace("\n", "</p><p>") + "</p>") if answer else "",
            "analysis": ("<p>" + html.escape(analysis).replace("\n", "</p><p>") + "</p>") if analysis else "",
            "knowledge": [],
            "category": subject,
            "source": source,
            "images": [],
        })
    return result


def _parse_text_columns(text: str, subject: str = "", grade: str = "", difficulty: int = 3,
                        source: str = "text"):
    """尝试按 CSV 列式模板解析（表头含 题型+题干）。非列式返回 None。"""
    import csv as _csv
    lines = [l for l in text.splitlines() if l.strip()]
    if not lines:
        return None
    first = lines[0]
    delim = None
    for d in ("\t", ",", ";"):
        if d in first and first.count(d) >= 2:
            delim = d
            break
    if delim is None:
        return None
    try:
        reader = _csv.reader(io.StringIO(text), delimiter=delim)
        rows = [r for r in reader if any((c or "").strip() for c in r)]
    except Exception:
        return None
    if not rows:
        return None
    head = [(h or "").replace(" ", "") for h in rows[0]]
    if not (any("题型" in h for h in head) and any("题干" in h for h in head)):
        return None
    idx = {}
    for j, h in enumerate(head):
        if "题型" in h: idx["type"] = j
        elif "题干" in h: idx["stem"] = j
        elif "选项" in h: idx["options"] = j
        elif "答案" in h: idx["answer"] = j
        elif "解析" in h: idx["analysis"] = j
        elif "难度" in h: idx["difficulty"] = j
        elif "学科" in h: idx["subject"] = j
        elif "年级" in h: idx["grade"] = j
        elif "知识点" in h: idx["knowledge"] = j
    if "type" not in idx or "stem" not in idx:
        return None
    qs, _, _ = _template_rows_to_questions(rows, 0, idx, None, difficulty, source, subject, grade)
    return qs


def _template_rows_to_questions(rows, header_idx, idx, img_by_row=None, difficulty=3,
                                source="xlsx", subject="", grade="", src_prefix="/images/"):
    """表格模板数据行 → 题目列表（xlsx / csv 共用）。

    rows: 全表 list[list[str]]；header_idx: 表头所在行号；idx: {字段: 列号}；
    img_by_row: {0-based行号: [图片文件名]}（xlsx 图片按行归属）或 None。
    返回 (questions, 首行学科, 首行年级)。
    """
    def _cell(r, key):
        j = idx.get(key)
        if j is None or j >= len(r):
            return ""
        return str(r[j] or "").strip()

    qe = _qe()
    type_cn = getattr(qe, "_TYPE_CN", {}) or {}
    out = []
    first_subject, first_grade = "", ""
    for i, r in enumerate(rows[header_idx + 1:]):
        if not any((c or "").strip() for c in r):
            continue
        stem = _cell(r, "stem")
        if not stem:
            continue
        typ_raw = _cell(r, "type")
        typ = type_cn.get(typ_raw) or qe.detect_type(stem, [])
        opts_raw = _cell(r, "options")
        opts = [o.strip() for o in re.split(r"[|；;，,]", opts_raw) if o.strip()]
        answer = _cell(r, "answer")
        analysis = _cell(r, "analysis")
        diff_raw = _cell(r, "difficulty")
        try:
            diff = int(diff_raw) if diff_raw else difficulty
        except ValueError:
            diff = difficulty
        if diff < 1 or diff > 5:
            diff = 3
        subj = _cell(r, "subject") or subject
        grd = _cell(r, "grade") or grade
        if not first_subject and subj:
            first_subject = subj
        if not first_grade and grd:
            first_grade = grd
        kp = [k.strip() for k in re.split(r"[|；;，,]", _cell(r, "knowledge")) if k.strip()]
        row_imgs = []
        ws_row = header_idx + 1 + i
        if img_by_row:
            row_imgs = img_by_row.get(ws_row, []) or []
        imgs_html = "".join(
            '<p><img src="%s" style="max-width:100%%;"></p>' % (src_prefix + f)
            for f in row_imgs)
        out.append({
            "id": "XL%03d" % (i + 1),
            "code": qe._gen_code(subj, grd, i + 1),
            "subject": subj,
            "grade": grd,
            "type": typ,
            "difficulty": diff,
            "score": 0,
            "stem": "<p>" + html.escape(stem).replace("\n", "</p><p>") + imgs_html + "</p>",
            "options": opts,
            "answer": ("<p>" + html.escape(answer).replace("\n", "</p><p>") + "</p>") if answer else "",
            "analysis": ("<p>" + html.escape(analysis).replace("\n", "</p><p>") + "</p>") if analysis else "",
            "knowledge": kp,
            "category": subj,
            "source": source,
            "images": row_imgs,
        })
    return out, first_subject, first_grade


def _xlsx_images_by_row(ws, img_dir):
    """提取工作表内图片，按锚点行（0-based）分组。返回 {row_idx: [fname]} 或 None。"""
    imgs = getattr(ws, "_images", None) or []
    if not imgs:
        return None
    os.makedirs(img_dir, exist_ok=True)
    out = {}
    for i, img in enumerate(imgs, 1):
        try:
            anchor = getattr(img, "anchor", None)
            row_idx = None
            if anchor is not None:
                fr = getattr(anchor, "from_", None) or getattr(anchor, "_from", None)
                if fr is not None:
                    row_idx = getattr(fr, "row", None)
            blob = None
            if hasattr(img, "_data") and callable(img._data):
                try:
                    blob = img._data()
                except Exception:
                    blob = None
            if not blob:
                ref = getattr(img, "ref", None)
                if isinstance(ref, str) and os.path.isfile(ref):
                    with open(ref, "rb") as f:
                        blob = f.read()
                elif isinstance(ref, (bytes, bytearray)):
                    blob = bytes(ref)
            if not blob:
                continue
            fname = "xlsx_img_%04d.png" % i
            with open(os.path.join(img_dir, fname), "wb") as f:
                f.write(blob)
            out.setdefault(row_idx, []).append(fname)
        except Exception:
            continue
    return out or None


def _parse_xlsx(data: bytes, filename: str, img_dir: str, subject: str = "", grade: str = "",
                difficulty: int = 3):
    """xlsx → 题目列表。优先表格模板列解析；无模板降级为编号行切分。
    返回 (questions, 模板学科, 模板年级)。"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    except Exception as e:
        raise ValidationError("Excel 解析失败（文件可能损坏或不是有效的 .xlsx）：%s" % e)
    try:
        ws = wb.active
        rows = []
        for row in ws.iter_rows(values_only=True):
            vals = [("" if v is None else str(v)).strip() for v in row]
            if any(vals):
                rows.append(vals)
        if not rows:
            return [], "", ""
        header_idx, idx = None, None
        for i, row in enumerate(rows):
            joined = "".join(row).replace(" ", "")
            if "题型" in joined and "题干" in joined:
                idx = {}
                for j, h in enumerate(row):
                    h2 = (h or "").replace(" ", "")
                    if "题型" in h2: idx["type"] = j
                    elif "题干" in h2: idx["stem"] = j
                    elif "选项" in h2: idx["options"] = j
                    elif "答案" in h2: idx["answer"] = j
                    elif "解析" in h2: idx["analysis"] = j
                    elif "难度" in h2: idx["difficulty"] = j
                    elif "学科" in h2: idx["subject"] = j
                    elif "年级" in h2: idx["grade"] = j
                    elif "知识点" in h2: idx["knowledge"] = j
                if "type" in idx and "stem" in idx:
                    header_idx = i
                    break
                idx = None
        img_by_row = _xlsx_images_by_row(ws, img_dir)
        if header_idx is not None:
            return _template_rows_to_questions(rows, header_idx, idx, img_by_row,
                                               difficulty, "xlsx", subject, grade)
        # 降级：每单元格一行 → 编号切分
        lines = []
        for r in rows:
            for c in r:
                if c.strip():
                    lines.append(c)
        qs = _parse_text("\n".join(lines), subject, grade, difficulty, source="xlsx")
        return qs, "", ""
    finally:
        try:
            wb.close()
        except Exception:
            pass


def _ocr_recognize(data: bytes, ext: str, subject: str, grade: str, difficulty: int,
                   url: str, api_key: str, model: str) -> list:
    """文件字节 → 视觉模型 OCR → 题目列表（含图题按 img_box 裁剪插入题干）。"""
    if ext == ".pdf":
        images = [("image/png", b) for b in _pdf_to_images(data)]
    else:
        mime = "image/png" if ext == ".png" else "image/jpeg"
        images = [(mime, _normalize_image(data))]
    prompt = _build_ocr_prompt(subject, grade)
    raw_list = []
    for i, (mime, img_bytes) in enumerate(images, 1):
        logger.info("smart_ocr_recognize_page", extra={"page": i, "total": len(images)})
        raw_list.append(_call_vision_model(url, api_key, model, img_bytes, mime, prompt))
    questions = []
    counter = 0
    for i, raw in enumerate(raw_list):
        mime, img_bytes = images[i] if i < len(images) else ("image/png", b"")
        for q in _parse_json_array(raw):
            counter += 1
            nq = _normalize_ocr_question(q, counter, subject, grade, difficulty)
            if not nq:
                continue
            box = q.get("img_box") or q.get("image_box") or None
            if box and img_bytes:
                try:
                    box = [int(x) for x in box]
                    if len(box) == 4 and box[0] < box[2] and box[1] < box[3]:
                        from PIL import Image as _PIL
                        pil = _PIL.open(io.BytesIO(img_bytes))
                        w, h = pil.size
                        x1 = max(0, box[0] - 6); y1 = max(0, box[1] - 6)
                        x2 = min(w, box[2] + 6); y2 = min(h, box[3] + 6)
                        if x2 - x1 > 10 and y2 - y1 > 10:
                            crop = pil.crop((x1, y1, x2, y2))
                            os.makedirs(DEMO_IMG_DIR, exist_ok=True)
                            fname = "ocr_img_%04d.png" % counter
                            crop.save(os.path.join(DEMO_IMG_DIR, fname), "PNG")
                            nq["stem"] = (nq.get("stem") or "") + (
                                '<p><img src="/images/%s" style="max-width:100%%;"></p>' % fname)
                            nq.setdefault("images", []).append(fname)
                except Exception as e:
                    logger.warning("ocr_crop_fail", extra={"err": str(e)})
            questions.append(nq)
    return questions


# 格式 key → 中文名（前端展示 / 日志）
_FORMAT_NAMES = {"docx": "Word", "pdf": "PDF", "png": "图片", "jpg": "图片", "xlsx": "Excel", "text": "文本"}


@router.post("/import-smart")
async def import_smart(
    file: UploadFile = File(...),
    difficulty: int = Form(3),
    subject: str = Form(""),
    grade: str = Form(""),
    provider: str = Form("zhipu"),
    api_key: str = Form(""),
    model: str = Form("glm-4v"),
    base_url: str = Form(""),
    _: Principal = Depends(require_permission("sync", "add")),
    db: Session = Depends(get_db),
):
    """智能导入：上传任意常见文件 → 自动识别格式 → 自动路由解析 → 智能推断学科/年级。

    支持：docx / pdf / png / jpg / xlsx / txt / csv / md（.doc/.xls 老格式给出明确提示）。
    返回 {code:0, data:{questions, count, format, detected_subject, detected_grade, img_questions, warnings}}。
    解析成功不落库，由前端预览确认后再调 POST /questions/import 入库。
    """
    filename = file.filename or "未命名文件"
    data = await file.read()
    if not data:
        return JSONResponse({"code": 400, "message": "上传文件为空，请重新选择文件", "data": None}, 400)
    if len(data) > _SMART_MAX_SIZE:
        return JSONResponse(
            {"code": 413, "message": "文件大小超过 30MB 上限，请拆分后分批导入", "data": None}, 413)

    fmt = _detect_file_format(data, filename)   # 识别不了会抛 ValidationError（友好提示）

    # ---------- OCR 类（pdf / png / jpg） ----------
    if fmt in ("pdf", "png", "jpg"):
        aicfg = get_ai_config(db)
        vcfg = aicfg.get("vision") or {}
        if "****" in (api_key or ""):   # 前端掩码回显值误传：视为未传
            api_key = ""
        if not (api_key or "").strip():
            if vcfg.get("api_key"):
                api_key = vcfg["api_key"]   # 数据库设置（env 已在 get_ai_config 内回退）
            else:
                raise ValidationError("图片/PDF 识别需要先配置视觉模型 API Key（系统设置→AI模型配置，或服务端 AI_ZHIPU_API_KEY）")
        model = (vcfg.get("model") or "").strip() or model
        url = (base_url or "").strip() or _VISION_PROVIDERS.get((provider or "").lower())
        if not url:
            raise ValidationError("不支持的视觉模型 provider：%s（可传 base_url 覆盖）" % provider)
        ext = ".pdf" if fmt == "pdf" else (".png" if fmt == "png" else ".jpg")
        questions = _ocr_recognize(data, ext, subject, grade, difficulty, url, api_key, model)
        if not questions:
            return JSONResponse(
                {"code": 422, "message": "未能识别出任何题目（请检查图片清晰度，或在系统设置中确认视觉模型配置正确）", "data": None}, 422)
        d_subject, d_grade = _smart_infer_meta(filename, "", subject, grade)
        img_questions = sum(1 for q in questions if "<img" in (q.get("stem") or ""))
        return {"code": 0, "message": "ok", "data": {
            "questions": questions, "count": len(questions), "format": fmt,
            "detected_subject": d_subject, "detected_grade": d_grade,
            "img_questions": img_questions, "warnings": [],
        }}

    warnings = []

    # ---------- docx（本地解析） ----------
    if fmt == "docx":
        qe = _qe()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        try:
            content_text = _extract_docx_text(tmp_path)
            d_subject, d_grade = _smart_infer_meta(filename, content_text, subject, grade)
            qs = qe.parse_docx(tmp_path, img_dir=DEMO_IMG_DIR,
                               subject=d_subject, grade=d_grade, difficulty=difficulty,
                               source="smart-docx", category=d_subject, src_prefix="/images/")
        finally:
            try:
                os.remove(tmp_path)
            except Exception:
                pass
        if not qs:
            return JSONResponse(
                {"code": 422, "message": "未能识别出任何题目：请确认文档按「1. 2. 3.」编号，或使用表格模板（点右上角「下载导出模板」）填写", "data": None}, 422)
        # 表格模板可能逐题带学科/年级，文件级推断结果取首题实际值
        d_subject = (qs[0].get("subject") or "") if qs else ""
        d_grade = (qs[0].get("grade") or "") if qs else ""
        img_questions = sum(1 for q in qs if "<img" in (q.get("stem") or ""))
        no_answer = sum(1 for q in qs if not (q.get("answer") or "").strip())
        if no_answer:
            warnings.append("%d 题缺少答案，请预览时核对" % no_answer)
        return {"code": 0, "message": "ok", "data": {
            "questions": qs, "count": len(qs), "format": fmt,
            "detected_subject": d_subject, "detected_grade": d_grade,
            "img_questions": img_questions, "warnings": warnings,
        }}

    # ---------- xlsx（本地解析） ----------
    if fmt == "xlsx":
        qs, tmpl_subj, tmpl_gr = _parse_xlsx(data, filename, DEMO_IMG_DIR, subject, grade, difficulty)
        d_subject, d_grade = _smart_infer_meta(filename, "", tmpl_subj or subject, tmpl_gr or grade)
        if not qs:
            return JSONResponse(
                {"code": 422, "message": "未能识别出任何题目：请确认表格包含「题型/题干」表头，或每行以「1. 2. 3.」编号开头", "data": None}, 422)
        for q in qs:
            if not q.get("subject"):
                q["subject"] = d_subject
            if not q.get("grade"):
                q["grade"] = d_grade
            if not q.get("category"):
                q["category"] = d_subject
        d_subject = (qs[0].get("subject") or "") if qs else ""
        d_grade = (qs[0].get("grade") or "") if qs else ""
        img_questions = sum(1 for q in qs if "<img" in (q.get("stem") or ""))
        no_answer = sum(1 for q in qs if not (q.get("answer") or "").strip())
        if no_answer:
            warnings.append("%d 题缺少答案，请预览时核对" % no_answer)
        return {"code": 0, "message": "ok", "data": {
            "questions": qs, "count": len(qs), "format": fmt,
            "detected_subject": d_subject, "detected_grade": d_grade,
            "img_questions": img_questions, "warnings": warnings,
        }}

    # ---------- 文本（txt / csv / md） ----------
    if fmt == "text":
        try:
            text = _decode_text(data)
        except ValueError as e:
            return JSONResponse({"code": 422, "message": "文本编码无法识别（%s）" % e, "data": None}, 422)
        # 扩展名伪装检测：扩展名是二进制格式但内容实为文本
        ext = os.path.splitext(filename)[1].lower()
        if ext in (".docx", ".pdf", ".png", ".jpg", ".jpeg", ".xlsx"):
            warnings.append("文件扩展名为 %s，但实际内容识别为文本，已按文本解析；若确实需要导入该二进制格式，请用对应软件重新导出" % ext)
        d_subject, d_grade = _smart_infer_meta(filename, text, subject, grade)
        qs = _parse_text(text, d_subject, d_grade, difficulty, source="smart-text")
        if not qs:
            return JSONResponse(
                {"code": 422, "message": "未能识别出任何题目：请确认文本按「1. 2. 3.」编号，每行一题，可含选项行与【答案】/【解析】标记",
                 "data": {"questions": [], "count": 0, "format": "text",
                          "detected_subject": d_subject, "detected_grade": d_grade,
                          "img_questions": 0, "warnings": warnings}}, 422)
        img_questions = 0
        no_answer = sum(1 for q in qs if not (q.get("answer") or "").strip())
        if no_answer:
            warnings.append("%d 题缺少答案，请预览时核对" % no_answer)
        return {"code": 0, "message": "ok", "data": {
            "questions": qs, "count": len(qs), "format": fmt,
            "detected_subject": d_subject, "detected_grade": d_grade,
            "img_questions": img_questions, "warnings": warnings,
        }}

    # 理论不可达（_detect_file_format 已兜底），防御性返回
    return JSONResponse({"code": 415, "message": "不支持的文件类型", "data": None}, 415)
