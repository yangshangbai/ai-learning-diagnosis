#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""试卷模板 & 答题卡模板 默认版生成引擎（python-docx）。

- generate_paper_template_docx：试卷 Word 母版（页眉信息 + 题型分节 + 连续题号 + 题干/选项/分值，图片回嵌）
- generate_sheet_template_docx：答题卡 Word 母版（机读带占位 + 四角对位 + 涂点/填空/方格稿纸分区）

设计要点：
1. 复用 question_import_export 的 _html_to_paras / _resolve_image，题干图片/公式不丢失
   （图片 src 兼容 相对路径 images/xxx.png、http(s) URL、data URI，逐一兜底）。
2. 不含答案（学生版）；generate_paper_template_docx 支持 with_answer=True 输出答案行（教师版）。
3. 文件命名由调用方（路由）负责：paper_<paperId>_<ts>.docx / sheet_<paperId>_<ts>.docx，
   本模块只负责把内容写进指定 out_path。
"""
import base64
import io
import json
import os
import re
import time
import urllib.request

import sys

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    raise SystemExit("请先安装 python-docx: pip install python-docx")

# 项目根的 question_import_export.py（解析引擎，含 _html_to_paras / _resolve_image）
BASE = os.path.dirname(os.path.abspath(__file__))  # backend/app
_IMPORT_EXPORT_DIR = os.path.dirname(os.path.dirname(BASE))  # 项目根目录
if _IMPORT_EXPORT_DIR not in sys.path:
    sys.path.insert(0, _IMPORT_EXPORT_DIR)
from question_import_export import _html_to_paras, _resolve_image  # noqa: E402

# 模板存储目录：生产优先 /opt/ai-learning/uploads/templates（nginx /uploads/ 指向），本地回退 backend/uploads/templates
_PRODUCTION_TEMPLATE_DIR = "/opt/ai-learning/uploads/templates"
_LOCAL_TEMPLATE_DIR = os.path.join(BASE, "uploads", "templates")

# 题型中文名
TYPE_CN = {
    "single_choice": "单选题",
    "multi_choice": "多选题",
    "true_false": "判断题",
    "fill_blank": "填空题",
    "essay": "解答题",
}

# 答题卡分区：选择 → 填空 → 解答
_SECTION_GROUPS = [
    ("choice", ("single_choice", "multi_choice", "true_false"), "选择题（涂点区）"),
    ("fill", ("fill_blank",), "填空题（填空框）"),
    ("essay", ("essay",), "解答题（方格稿纸）"),
]


def get_template_dir() -> str:
    """返回模板存储目录（自动建目录）。"""
    d = _PRODUCTION_TEMPLATE_DIR if os.path.isdir("/opt/ai-learning") else _LOCAL_TEMPLATE_DIR
    os.makedirs(d, exist_ok=True)
    return d


def _coerce(v, default=""):
    if v is None:
        return default
    return str(v)


# ---------------------------------------------------------------- 图片兜底解析
def _resolve_image_enhanced(src: str, tmp_dir: str):
    """增强版图片解析：先走 question_import_export 的本地路径候选，再兜底 http(s) / data URI。"""
    if not src:
        return None
    # 1) 本地候选（demo/images 等）
    path = _resolve_image(src)
    if path and os.path.isfile(path):
        return path
    # 2) http(s) URL：下载到临时目录
    if src.startswith("http://") or src.startswith("https://"):
        try:
            os.makedirs(tmp_dir, exist_ok=True)
            name = os.path.basename(src.split("?")[0]) or ("img_%s" % int(time.time() * 1000))
            if not re.search(r"\.(png|jpe?g|gif|bmp|svg)$", name, re.I):
                name = "img_%s.png" % int(time.time() * 1000)
            fpath = os.path.join(tmp_dir, name)
            urllib.request.urlretrieve(src, fpath)
            if os.path.isfile(fpath):
                return fpath
        except Exception:
            return None
    # 3) data URI：base64 解码落盘
    if src.startswith("data:"):
        try:
            m = re.match(r"data:image/(png|jpe?g|gif|bmp|svg\+xml);base64,(.+)", src, re.S)
            if m:
                os.makedirs(tmp_dir, exist_ok=True)
                ext = {"png": "png", "jpg": "jpg", "jpeg": "jpg", "gif": "gif", "bmp": "bmp", "svg+xml": "svg"}.get(m.group(1), "png")
                fpath = os.path.join(tmp_dir, "img_%s.%s" % (int(time.time() * 1000), ext))
                with open(fpath, "wb") as f:
                    f.write(base64.b64decode(m.group(2)))
                return fpath
        except Exception:
            return None
    return None


def _html_to_paras_enhanced(doc, html_str, tmp_dir):
    """包装 question_import_export._html_to_paras：临时替换其模块级 _resolve_image 为增强版。"""
    from question_import_export import _html_to_paras as _orig_html_to_paras

    import question_import_export as qe

    old = qe._resolve_image
    qe._resolve_image = lambda src: _resolve_image_enhanced(src, tmp_dir)
    try:
        return _orig_html_to_paras(doc, html_str)
    finally:
        qe._resolve_image = old


def _set_run_font(run, size_pt=11, bold=False, cn_font="宋体"):
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = cn_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)


def _add_par(doc, text="", size_pt=11, bold=False, align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        _set_run_font(run, size_pt, bold)
    return p


# ---------------------------------------------------------------- 分组与编号
def _group_questions(questions):
    """按 选择→填空→解答 分组（组内保持原顺序），全卷连续编号。
    返回 [ {group_title, seq_title, items:[{...item, no}]} ]，item 含 type/stem/options/answer/score。"""
    result = []
    seq = 0
    cn_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    for gi, (key, types, title) in enumerate(_SECTION_GROUPS):
        items = [q for q in questions if q.get("type") in types]
        if not items:
            continue
        for it in items:
            seq += 1
            it["no"] = seq
        result.append(
            {
                "key": key,
                "group_title": "%s、%s" % (cn_nums[gi], title),
                "items": items,
            }
        )
    return result


def _paper_header_pars(doc, paper):
    """页眉：试卷编号/学科/年级/分类/题数/总分。"""
    h = doc.sections[0].header
    p = h.paragraphs[0] if h.paragraphs else h.add_paragraph()
    p.text = ""
    info = "试卷编号：%s    学科：%s    年级：%s    分类：%s    题数：%s    总分：%s" % (
        _coerce(paper.get("paper_code")),
        _coerce(paper.get("subject")),
        _coerce(paper.get("grade")),
        _coerce(paper.get("category")),
        paper.get("question_count") or 0,
        paper.get("total_score") or 0,
    )
    run = p.add_run(info)
    _set_run_font(run, 9, False, "宋体")


def _section_title(doc, title, total_score=None):
    p = _add_par(doc, title, size_pt=12, bold=True, space_after=6)
    return p


def _add_question_stem(doc, item, tmp_dir):
    """题干行：'{no}、题干HTML（含分值）'，图片/公式回嵌。"""
    no = item.get("no")
    score = item.get("score") or 0
    head = "%d、" % no if no else ""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if head:
        run = p.add_run(head)
        _set_run_font(run, 11, True)
    stem = _coerce(item.get("stem"))
    _html_to_paras_enhanced(doc, stem, tmp_dir) if stem else None
    if score:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(4)
        sr = sp.add_run("（%d 分）" % score)
        _set_run_font(sr, 9, False, "宋体")


# ---------------------------------------------------------------- 试卷模板
def generate_paper_template_docx(paper: dict, questions: list, out_path: str, with_answer: bool = False) -> str:
    """生成试卷 Word 模板（学生版默认不含答案；with_answer=True 附答案行）。

    :param paper: {"paper_code","name","subject","grade","category","question_count","total_score"}
    :param questions: [{"type","stem","options":[...],"answer","score"}]
    :param out_path: 输出 .docx 路径
    :return: out_path
    """
    tmp_dir = os.path.join(os.path.dirname(out_path), "_img_cache")
    doc = Document()
    # 默认样式：中文字体
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)

    _paper_header_pars(doc, paper)

    # 标题
    _add_par(doc, _coerce(paper.get("name"), "试卷"), size_pt=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    groups = _group_questions(questions)
    if not groups:
        # 不应走到这里（路由已校验有题目）；防御性返回空文档
        doc.save(out_path)
        return out_path

    for g in groups:
        _section_title(doc, g["group_title"] + "（共 %d 题）" % len(g["items"]))
        for item in g["items"]:
            _add_question_stem(doc, item, tmp_dir)
            opts = item.get("options") or []
            if item["type"] in ("single_choice", "multi_choice") and opts:
                for o in opts:
                    op = doc.add_paragraph()
                    op.paragraph_format.left_indent = Cm(0.75)
                    op.paragraph_format.space_after = Pt(1)
                    orun = op.add_run("  " + _coerce(o))
                    _set_run_font(orun, 11, False)
            elif item["type"] == "true_false":
                op = doc.add_paragraph()
                op.paragraph_format.left_indent = Cm(0.75)
                op.paragraph_format.space_after = Pt(1)
                orun = op.add_run("  （  ）正确    （  ）错误")
                _set_run_font(orun, 11, False)
            if with_answer and item.get("answer"):
                ap = doc.add_paragraph()
                ap.paragraph_format.space_after = Pt(4)
                ar = ap.add_run("【答案】")
                _set_run_font(ar, 10, True)
                _html_to_paras_enhanced(doc, _coerce(item.get("answer")), tmp_dir)

    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------- 答题卡模板
def _choice_row(doc, item):
    """选择涂点行：单选○ / 多选□ / 判断√×。"""
    no = item.get("no")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("%d、 " % no)
    _set_run_font(run, 11, True)
    typ = item.get("type")
    if typ == "multi_choice":
        opts = item.get("options") or []
        for o in opts:
            letter = _coerce(o).strip()[:1] or "?"
            r = p.add_run("□ %s   " % letter)
            _set_run_font(r, 11, False)
    elif typ == "true_false":
        for t in ("√", "×"):
            r = p.add_run("○ %s   " % t)
            _set_run_font(r, 11, False)
    else:  # single_choice
        opts = item.get("options") or []
        for o in opts:
            letter = _coerce(o).strip()[:1] or "?"
            r = p.add_run("○ %s   " % letter)
            _set_run_font(r, 11, False)


def _blank_row(doc, item):
    """填空框：连续下划线。"""
    no = item.get("no")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("%d、 " % no)
    _set_run_font(run, 11, True)
    r = p.add_run("＿＿＿＿＿＿          ＿＿＿＿＿＿")
    _set_run_font(r, 11, False)


def _essay_grid(doc, item):
    """解答方格稿纸：表格模拟方格，高度≈score*8。"""
    no = item.get("no")
    score = item.get("score") or 0
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("%d、 " % no)
    _set_run_font(run, 11, True)
    # 前端 min-height = score*8(px)；Word 表格格 0.6cm，行数≈高度/0.6cm
    rows = max(3, int(round((score * 8) / 18)))
    cols = 25
    table = doc.add_table(rows=rows, cols=cols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(0.62)
            cell.height = Cm(0.62)
    # 表后留白
    _add_par(doc, "", size_pt=4, space_after=4)


def generate_sheet_template_docx(paper: dict, questions: list, out_path: str) -> str:
    """生成答题卡 Word 模板：机读带 + 四角对位 + 涂点/填空/方格稿纸分区。"""
    tmp_dir = os.path.join(os.path.dirname(out_path), "_img_cache")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)

    # ---- 页眉机读带（含 QR 占位文本，任务级渲染时按 payload 替换）----
    h = doc.sections[0].header
    hp = h.paragraphs[0] if h.paragraphs else h.add_paragraph()
    hp.text = ""
    info = "任务：____________    试卷：%s    学生：____________" % _coerce(paper.get("paper_code") or paper.get("name"))
    run = hp.add_run(info)
    _set_run_font(run, 9, False)
    qr_text = "[QR:JY|tk={task}|pp={paper}|st={student}|sb={subject}|pg=1/1]"
    qp = h.add_paragraph()
    qp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    qr = qp.add_run(qr_text)
    _set_run_font(qr, 8, False, "宋体")

    # ---- 四角对位标记 ----
    top = doc.add_paragraph()
    top.paragraph_format.space_after = Pt(6)
    tr = top.add_run("┌")
    _set_run_font(tr, 14, True)
    top.add_run("            ")
    _set_run_font(top.add_run("┐"), 14, True)

    # 标题
    _add_par(doc, "%s 答题卡" % _coerce(paper.get("name"), "答题卡"), size_pt=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    groups = _group_questions(questions)
    for g in groups:
        _section_title(doc, g["group_title"] + "（共 %d 题）" % len(g["items"]))
        for item in g["items"]:
            if g["key"] == "choice":
                _choice_row(doc, item)
            elif g["key"] == "fill":
                _blank_row(doc, item)
            else:
                _essay_grid(doc, item)

    # 四角对位标记（下）
    bottom = doc.add_paragraph()
    bottom.paragraph_format.space_before = Pt(12)
    br = bottom.add_run("└")
    _set_run_font(br, 14, True)
    bottom.add_run("            ")
    _set_run_font(bottom.add_run("┘"), 14, True)

    # 页脚对位（文档页脚也放一组）
    ft = doc.sections[0].footer
    fp = ft.paragraphs[0] if ft.paragraphs else ft.add_paragraph()
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fl = fp.add_run("└")
    _set_run_font(fl, 10, True)
    fp.add_run("                        ")
    _set_run_font(fp.add_run("┘"), 10, True)

    doc.save(out_path)
    return out_path


def layout_from_questions(questions: list) -> dict:
    """由题目列表生成 layout_config（与 paper.py _build_layout 同构；引擎自测/兜底用）。

    :param questions: [{"type","score"}]（按组卷顺序）
    """
    choice, fill, essay = [], [], []
    for idx, q in enumerate(questions, 1):
        typ = q.get("type") or "essay"
        score = q.get("score") or 0
        entry = {"question_number": idx, "score": score}
        if typ in ("single_choice", "multi_choice", "true_false"):
            w = 10 if typ == "true_false" else 15
            choice.append({**entry, "width_mm": w})
        elif typ == "fill_blank":
            fill.append({**entry, "width_mm": 40})
        else:
            essay.append({**entry, "height_mm": max(40, (score or 8) * 10)})
    sections = []
    if choice:
        sections.append({"type": "choice", "questions": choice, "width_mm": sum(c["width_mm"] for c in choice)})
    if fill:
        sections.append({"type": "fill", "questions": fill, "width_mm": sum(f["width_mm"] for f in fill)})
    if essay:
        sections.append({"type": "essay", "questions": essay, "height_mm": sum(e["height_mm"] for e in essay)})
    return {"page_size": "A4", "sections": sections}


# ---------------------------------------------------------------------------
# 任务级每生答题卡 Word（2026-08-30 需求：每生一页 + 个人二维码，payload 含 任务/试卷/班级/学生）
# ---------------------------------------------------------------------------
def build_sheet_qr_payload(task: dict, paper: dict, student: dict) -> str:
    """每生二维码 payload：任务号/试卷号/班级号/学生号 + 页码。"""
    return "JY|tk={}|pp={}|cl={}|st={}|pg=1/1".format(
        (task or {}).get("code") or "",
        (paper or {}).get("code") or "",
        (student or {}).get("classCode") or "",
        (student or {}).get("code") or "",
    )


def _qr_png_stream(payload: str):
    """二维码 PNG 流（与 answer_sheet_renderer.generate_qr_data_url 同规格）。"""
    try:
        import qrcode

        img = qrcode.make(payload)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf
    except Exception:
        return None


def generate_task_sheets_docx(task: dict, paper: dict, students: list, questions: list, out_path: str) -> str:
    """任务级每生答题卡 Word：每生一页，页顶机读带含真实个人二维码图片。

    payload = JY|tk=任务号|pp=试卷号|cl=班级号|st=学生号|pg=1/1
    版式为系统默认（用户自定义模板与每生 QR 的占位合并列为后续迭代）。
    """
    from docx.shared import Inches

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)

    groups = _group_questions(questions)
    stu_list = students or [None]
    for idx, stu in enumerate(stu_list):
        if idx > 0:
            doc.add_page_break()

        # ---- 机读带（表格：左信息 / 右二维码图片）----
        payload = build_sheet_qr_payload(task, paper, stu or {})
        tbl = doc.add_table(rows=1, cols=2)
        left, right = tbl.rows[0].cells
        code = _coerce((stu or {}).get("code"))
        sname = _coerce((stu or {}).get("name"))
        cname = _coerce((stu or {}).get("className"))
        ccode = _coerce((stu or {}).get("classCode"))
        p1 = left.paragraphs[0]
        r1 = p1.add_run("任务：%s %s" % (_coerce((task or {}).get("code")), _coerce((task or {}).get("name"))))
        _set_run_font(r1, 9, False)
        for line in (
            "试卷：%s" % _coerce((paper or {}).get("code") or (paper or {}).get("name")),
            "班级：%s %s" % (ccode, cname) if (ccode or cname) else "班级：____________",
            "学生：%s %s" % (code, sname) if (code or sname) else "学生：____________",
        ):
            rp = left.add_paragraph()
            rr = rp.add_run(line)
            _set_run_font(rr, 9, False)
        stream = _qr_png_stream(payload)
        if stream:
            right.paragraphs[0].add_run().add_picture(stream, width=Inches(1.1))
            cap = right.add_paragraph()
            cr = cap.add_run("扫码识别")
            _set_run_font(cr, 8, False)

        # ---- 四角对位（上）----
        top = doc.add_paragraph()
        top.paragraph_format.space_after = Pt(6)
        tr = top.add_run("┌")
        _set_run_font(tr, 14, True)
        top.add_run("            ")
        _set_run_font(top.add_run("┐"), 14, True)

        # ---- 标题 ----
        title = "%s 答题卡" % _coerce((paper or {}).get("name"), "答题卡")
        if code or sname:
            title += "（%s %s）" % (code, sname)
        _add_par(doc, title, size_pt=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

        # ---- 三分区 ----
        for g in groups:
            _section_title(doc, g["group_title"] + "（共 %d 题）" % len(g["items"]))
            for item in g["items"]:
                if g["key"] == "choice":
                    _choice_row(doc, item)
                elif g["key"] == "fill":
                    _blank_row(doc, item)
                else:
                    _essay_grid(doc, item)

        # ---- 四角对位（下）----
        bottom = doc.add_paragraph()
        bottom.paragraph_format.space_before = Pt(12)
        br = bottom.add_run("└")
        _set_run_font(br, 14, True)
        bottom.add_run("            ")
        _set_run_font(bottom.add_run("┘"), 14, True)

    doc.save(out_path)
    return out_path


def generate_merged_sheets_zip(task: dict, paper: dict, students: list, template_path: str, out_zip: str) -> str:
    """用户自定义模板(source=user) × 每生二维码合并：每生一份 docx 打包为 zip。

    注入规则（对页眉/正文/表格全部段落生效，幂等）：
      - 含 [QR:...] 占位的段落 → 剥离占位文本并在段末追加该生二维码图片
      - '任务：_+' → 任务号+名称；'学生：_+' → 学生号+姓名（班级）
    未含占位的模板仅做任务/学生信息回填（二维码追加到最后一个非空段落）。
    """
    import shutil
    import zipfile

    tk, tname = _coerce((task or {}).get("code")), _coerce((task or {}).get("name"))
    pp = _coerce((paper or {}).get("code") or (paper or {}).get("name"))
    tmp_dir = os.path.join(os.path.dirname(out_zip) or ".", "_merged_sheets")
    os.makedirs(tmp_dir, exist_ok=True)
    files = []
    for idx, stu in enumerate(students, 1):
        code = _coerce((stu or {}).get("code")) or ("S%02d" % idx)
        sname = _coerce((stu or {}).get("name"))
        cls = _coerce((stu or {}).get("classCode"))
        payload = build_sheet_qr_payload(task, paper, stu or {})
        stream = _qr_png_stream(payload)

        doc = Document(template_path)

        def _walk(paragraphs):
            injected = False
            for para in paragraphs:
                full = "".join(r.text for r in para.runs)
                if not full.strip():
                    continue
                new = re.sub(r"\[QR:[^\]]*\]", "", full)
                new = re.sub(r"任务：_+", "任务：%s %s" % (tk, tname), new)
                new = re.sub(r"学生：_+", "学生：%s %s（%s）" % (code, sname, cls), new)
                had_qr = "[QR:" in full
                if new == full and not had_qr:
                    continue
                for r in para.runs:
                    r.text = ""
                anchor = para.runs[0] if para.runs else para.add_run("")
                anchor.text = new
                if had_qr and stream is not None:
                    stream.seek(0)
                    para.add_run().add_picture(stream, width=Inches(1.0))
                    injected = True
            return injected

        _walk(doc.sections[0].header.paragraphs)
        injected_body = _walk(doc.paragraphs)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if _walk(cell.paragraphs):
                        injected_body = True
        if not injected_body and stream is not None:
            # 模板未含 QR 占位：在正文末尾追加机读带段落（信息+二维码）
            stream.seek(0)
            tail = doc.add_paragraph()
            tr = tail.add_run("任务：%s %s  试卷：%s  学生：%s %s  " % (tk, tname, pp, code, sname))
            _set_run_font(tr, 9, False)
            tail.add_run().add_picture(stream, width=Inches(1.0))

        fpath = os.path.join(tmp_dir, "%s-答题卡-%s.docx" % (tk or "TASK", code))
        doc.save(fpath)
        files.append((fpath, "%s-答题卡-%s-%s.docx" % (tk or "TASK", code, sname)))

    with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as z:
        for fpath, arcname in files:
            z.write(fpath, arcname)
    shutil.rmtree(tmp_dir, ignore_errors=True)
    return out_zip
