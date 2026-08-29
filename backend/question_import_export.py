#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""题库导入 / 导出：Word(.docx) ↔ 系统题目 JSON（兼容全学科）

系统题目规范格式（与 demo/questions_data.js 的 REAL_QUESTIONS、后端 Question 模型一致）：
    {
      "id": "RQ01", "code": "MAT-G7-KP420-0001",
      "subject": "数学", "grade": "初一",
      "type": "single_choice|multi_choice|fill_blank|true_false|essay",
      "difficulty": 1~5, "score": 数值,
      "stem": "<p>题干 HTML（含 <img> / MathJax <svg>）</p>",
      "options": ["A. ...", "B. ..."],   # 选择题才有
      "answer": "<p>答案 HTML</p>",
      "analysis": "<p>解析 HTML</p>",
      "knowledge": ["知识点"], "category": "分类", "source": "manual|docx"
    }

本模块能力：
  1. check_docx(path)                 —— 格式检查（题数/图片数/是否含答案解析/题型分布）
  2. parse_docx(path, img_dir, ...)   —— docx → 题目列表（题干含图片提取为文件并内联 <img>）
  3. export_docx(questions, path, ...) —— 题目列表 → docx（题目+【答案】+【解析】，图片回嵌）
  4. detect_type(blocks) / classify     —— 题型识别

依赖：python-docx（pip install python-docx）
"""
import html
import os
import re
import json
import uuid

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
except ImportError:
    raise SystemExit("请先安装 python-docx: pip install python-docx")

IMG_EXT = {"image/png": "png", "image/jpeg": "jpg", "image/gif": "gif", "image/bmp": "bmp", "image/svg+xml": "svg"}

# 学科/年级 → 位置编码码段
SUBJECT_CODES = {"数学": "MAT", "语文": "CHN", "英语": "ENG", "物理": "PHY", "化学": "CHM",
                 "生物": "BIO", "历史": "HIS", "地理": "GEO", "政治": "POL", "道德与法治": "POL",
                 "科学": "SCI"}
GRADE_CODES = {"一年级": "G1", "二年级": "G2", "三年级": "G3", "四年级": "G4", "五年级": "G5",
               "六年级": "G6", "初一": "G7", "初二": "G8", "初三": "G9", "高一": "G10",
               "高二": "G11", "高三": "G12"}


def _gen_code(subject, grade, idx):
    """生成位置编码：学科码-年级码-IMP-序号4位（无知识点信息时用 IMP 占位）。"""
    sc = SUBJECT_CODES.get(subject, "IMP")
    gc = GRADE_CODES.get(grade, "G1")
    return "%s-%s-IMP-%04d" % (sc, gc, idx)


# ---------------------------------------------------------------- 文档块读取
def iter_blocks(doc):
    """按文档顺序遍历 body，产出块：('para', html, [images]) 或 ('table', rows)。"""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("para", child)
        elif child.tag == qn("w:tbl"):
            yield ("table", child)


def _para_html(p_el, doc, img_dir, counter, src_prefix=""):
    """把单个段落 XML 转成 HTML 片段（文本 + 内联图片按出现顺序），图片落到 img_dir。"""
    out = []
    img_files = []
    # 每个解析会话生成唯一前缀（counter[1]），同一文件多图共享前缀，不同文件互不覆盖
    session_tag = counter[1]
    # 依序遍历段落内所有 w:t 与 a:blip，保留图文顺序
    for node in p_el.iter():
        if node.tag == qn("w:t"):
            t = node.text or ""
            if t:
                out.append(html.escape(t))
        elif node.tag == qn("a:blip"):
            rid = node.get(qn("r:embed"))
            if not rid:
                continue
            try:
                part = doc.part.related_parts[rid]
            except Exception:
                continue
            ext = IMG_EXT.get(part.content_type, "png")
            os.makedirs(img_dir, exist_ok=True)
            fname = "img_%s_%04d.%s" % (session_tag, counter[0], ext)
            counter[0] += 1
            with open(os.path.join(img_dir, fname), "wb") as f:
                f.write(part.blob)
            img_files.append(fname)
            out.append('<img src="%s">' % (src_prefix + fname))
    txt = "".join(out).strip()
    return (txt, img_files)


def _table_rows(tbl_el):
    rows = []
    for tr in tbl_el.iter(qn("w:tr")):
        cells = []
        for tc in tr.iter(qn("w:tc")):
            c = "".join(n.text or "" for n in tc.iter() if n.tag == qn("w:t")).strip()
            cells.append(c)
        rows.append(cells)
    return rows


def _table_cell_html(tc_el, doc, img_dir, counter, src_prefix=""):
    """把表格单元格内容转 HTML（文本 + 单元格内段落的图片），图片按出现顺序落盘。"""
    out = []
    imgs = []
    for p_el in tc_el.iter(qn("w:p")):
        h, imgs_p = _para_html(p_el, doc, img_dir, counter, src_prefix)
        if h:
            out.append(h)
        imgs.extend(imgs_p)
    return "<br>".join(out), imgs


def _read_doc_blocks(path, img_dir, src_prefix=""):
    doc = Document(path)
    counter = [0, uuid.uuid4().hex[:6]]  # [图片序号, 会话唯一前缀]
    blocks = []  # ('para', (html, imgs)) | ('table', (html, imgs))  —— 统一为 (html, imgs)，图片不丢
    for kind, el in iter_blocks(doc):
        if kind == "para":
            h, imgs = _para_html(el, doc, img_dir, counter, src_prefix)
            blocks.append(("para", (h, imgs)))
        else:
            # 表格：逐单元格转 HTML（含单元格内段落的图片），避免表格排版 docx 丢图
            h, imgs = _table_html(el, doc, img_dir, counter, src_prefix)
            blocks.append(("table", (h, imgs)))
    return doc, blocks, counter[0]


def _table_html(tbl_el, doc, img_dir, counter, src_prefix=""):
    """表格 → HTML（含单元格内图片）。返回 (html, imgs)。"""
    rows = []
    all_imgs = []
    for tr in tbl_el.iter(qn("w:tr")):
        cells = []
        for tc in tr.iter(qn("w:tc")):
            h, imgs = _table_cell_html(tc, doc, img_dir, counter, src_prefix)
            cells.append(h)
            all_imgs.extend(imgs)
        rows.append("<tr>" + "".join("<td>%s</td>" % c for c in cells) + "</tr>")
    return "<table border='1'>" + "".join(rows) + "</table>", all_imgs


# ---------------------------------------------------------------- 题型识别
def detect_type(text, opts):
    """根据文本特征判断题型。opts 为选项列表。"""
    t = text or ""
    has_opt = bool(opts) or bool(re.search(r"(^|[\s\n])([A-D])[\.、．]", t))
    if re.search(r"(多选|不定项)", t):
        return "multi_choice"
    if has_opt and re.search(r"(正确|错误|对错|判断题|√|×)", t):
        return "true_false"
    if has_opt:
        return "single_choice"
    if re.search(r"(判断|正确|错误)", t) and not has_opt:
        return "true_false"
    # 填空：出现连续 ≥2 个下划线即判定（单个下划线如"__"也判；需 ≥1 处）
    if len(re.findall(r"[＿_]{2,}", t)) >= 1 or re.search(r"填空", t):
        return "fill_blank"
    return "essay"


# ---------------------------------------------------------------- 题目切分
def _split_questions(blocks, default_subject="", default_grade=""):
    """按 1. 2. 3. 编号把块切分成题目。返回 list of {title, stem_blocks, answer_blocks, analysis_blocks}。"""
    # 先找编号段落位置
    num_re = re.compile(r"^\s*(\d{1,3})\s*[\.、．]\s*")
    questions = []
    cur = None
    pending_head = []  # 编号之前的块（试卷标题等）

    def flush(cur):
        if cur and (cur["stem"] or cur["answer"]):
            questions.append(cur)

    for kind, payload in blocks:
        if kind == "para":
            m = num_re.match(payload[0])
            if m:
                flush(cur)
                cur = {"num": int(m.group(1)), "title": payload[0],
                       "stem": [], "answer": [], "analysis": [], "imgs": []}
                # 编号行其余内容也算题干
                rest = payload[0][m.end():].strip()
                if rest:
                    cur["stem"].append(rest)
                cur["imgs"] += payload[1]
                continue
            elif payload[0].strip().startswith("【答案】") or payload[0].strip().startswith("[答案]"):
                if cur is not None:
                    cur["answer"].append(payload[0].strip())
                continue
            elif payload[0].strip().startswith("【解析】") or payload[0].strip().startswith("[解析]"):
                if cur is not None:
                    cur["analysis"].append(payload[0].strip())
                continue
            if cur is None:
                pending_head.append(payload)
            else:
                # 属于当前题的题干（未遇到答案/解析标记前）
                if not cur["answer"] and not cur["analysis"]:
                    cur["stem"].append(payload[0])
                    cur["imgs"] += payload[1]
                elif cur["analysis"]:
                    cur["analysis"].append(payload[0])
                else:
                    cur["answer"].append(payload[0])
        else:  # table（payload 为 (html, imgs)，与 para 一致）
            if cur is None:
                pending_head.append(("para", payload))
            elif not cur["answer"] and not cur["analysis"]:
                cur["stem"].append(payload[0])
                cur["imgs"] += payload[1]
            elif cur["analysis"]:
                cur["analysis"].append(payload[0])
                cur["imgs"] += payload[1]
            else:
                cur["answer"].append(payload[0])
                cur["imgs"] += payload[1]
    flush(cur)
    return questions


def _table_to_html(rows):
    if not rows:
        return ""
    h = "<table border='1'>"
    for r in rows:
        h += "<tr>" + "".join("<td>%s</td>" % html.escape(c) for c in r) + "</tr>"
    return h + "</table>"


# ---------------------------------------------------------------- 格式检查
def check_docx(path, img_dir=".import_imgs"):
    """格式检查，返回报告 dict。"""
    doc, blocks, n_img = _read_doc_blocks(path, img_dir)
    questions = _split_questions(blocks)
    report = {
        "file": path,
        "ok": len(questions) > 0,
        "questions": len(questions),
        "images": n_img,
        "has_answer": sum(1 for q in questions if q["answer"]),
        "has_analysis": sum(1 for q in questions if q["analysis"]),
        "types": {},
        "messages": [],
    }
    if not report["ok"]:
        report["messages"].append("未识别到任何题目（需以「1. 2. 3.」编号）")
    for q in questions:
        t = detect_type(" ".join(q["stem"]), None)
        report["types"][t] = report["types"].get(t, 0) + 1
    if report["has_answer"] == 0:
        report["messages"].append("未检测到【答案】标记（教师版应含答案）")
    return report


# ---------------------------------------------------------------- 导入
def parse_docx(path, img_dir=".import_imgs", subject="", grade="", difficulty=3,
               source="docx", category="", src_prefix=""):
    """docx → 题目列表（系统 JSON 格式）。"""
    _, blocks, _ = _read_doc_blocks(path, img_dir, src_prefix)
    # 优先识别表格模板（表头含 题型+题干）：人手友好的导入模板
    tpl = _parse_table_template(blocks, default_subject=subject, default_grade=grade,
                                default_difficulty=difficulty, source=source, category=category,
                                src_prefix=src_prefix)
    if tpl is not None:
        return tpl
    questions = _split_questions(blocks, subject, grade)
    result = []
    for i, q in enumerate(questions, 1):
        stem_html = "<p>" + "</p><p>".join(x for x in q["stem"] if x) + "</p>"
        answer_html = "<p>" + "</p><p>".join(x.lstrip("【答案】").lstrip("[答案]").strip() for x in q["answer"] if x) + "</p>"
        analysis_html = "<p>" + "</p><p>".join(x.lstrip("【解析】").lstrip("[解析]").strip() for x in q["analysis"] if x) + "</p>"
        stem_text = " ".join(q["stem"])
        opts = re.findall(r"(?:^|[\s\n])([A-D])[\.、．]\s*", stem_text)
        typ = detect_type(stem_text, opts)
        result.append({
            "id": "IMP%03d" % q["num"],
            "code": _gen_code(subject, grade, i),
            "subject": subject,
            "grade": grade,
            "type": typ,
            "difficulty": difficulty,
            "score": 0,
            "stem": stem_html,
            "options": list(dict.fromkeys(opts)),
            "answer": answer_html,
            "analysis": analysis_html,
            "knowledge": [],
            "category": category,
            "source": source,
            "images": q["imgs"],
        })
    return result


# ---------------------------------------------------------------- 导出
BASE = os.path.dirname(os.path.abspath(__file__))

_LATEX_MAP = [
    (r"\\angle", "∠"), (r"\\circ", "°"), (r"\\times", "×"), (r"\\div", "÷"),
    (r"\\cdot", "·"), (r"\\leq", "≤"), (r"\\geq", "≥"), (r"\\neq", "≠"),
    (r"\\pm", "±"), (r"\\infty", "∞"), (r"\\triangle", "△"), (r"\\parallel", "∥"),
    (r"\\perp", "⊥"), (r"\\angle", "∠"), (r"\\frac\s*\{\s*([^{}]*?)\s*\}\s*\{\s*([^{}]*?)\s*\}", r"(\1)/(\2)"),
    (r"\\sqrt\s*\{\s*([^{}]*?)\s*\}", r"√(\1)"),
    (r"\^{([^{}]+)}", r"^\1"), (r"_{([^{}]+)}", r"_\1"),
]


def _latex_to_text(s):
    """把 MathJax SVG 的 <title> LaTeX 转成可读文本（避免公式丢失）。"""
    s = html.unescape(s or "")
    for pat, rep in _LATEX_MAP:
        s = re.sub(pat, rep, s)
    s = s.replace("{", "").replace("}", "").replace("\\", "")
    return s.strip()


def _resolve_image(src, img_dir=".import_imgs"):
    """按优先级解析 <img src> 到本地文件（嵌入 Word 用，不依赖 URL）。"""
    if not src:
        return None
    candidates = []
    if os.path.isabs(src):
        candidates.append(src)
    candidates.append(os.path.join(img_dir, src))
    candidates.append(os.path.join(BASE, "demo", src))          # demo/images/xxx.png
    candidates.append(os.path.join(BASE, src))
    candidates.append(os.path.join(BASE, "demo", "images", os.path.basename(src)))
    for c in candidates:
        if os.path.isfile(c):
            return c
    return None


def _html_to_paras(doc, html_str):
    """把 HTML 拆成段落并写入 doc：<img> 嵌入本地图片、<svg> 公式转文本、其余转文本。"""
    html_str = _coerce_str(html_str)
    paras = []
    # 精确匹配 <p>/<p ...>/</p> 段落标签，避免误切 SVG 内的 <path>/<polygon> 等
    segs = re.split(r"</?p(?:\s[^>]*)?>", html_str)
    for seg in segs:
        seg = seg.strip()
        if not seg:
            continue
        p = doc.add_paragraph()
        # 依序处理：<img> / <svg>公式 / 纯文本
        pos = 0
        for m in re.finditer(r'<img\s+[^>]*src="([^"]+)"[^>]*>|<svg[^>]*>.*?</svg>', seg, re.S):
            # 前导文本
            txt = re.sub(r"<[^>]+>", "", seg[pos:m.start()])
            txt = html.unescape(txt)
            if txt.strip():
                p.add_run(txt)
            token = m.group(0)
            if token.startswith("<img"):
                src = m.group(1)
                path = _resolve_image(src)
                if path:
                    try:
                        p.add_run().add_picture(path, width=None)
                    except Exception:
                        p.add_run("[图]")
                else:
                    p.add_run("[图]")
            else:  # svg 公式 → 提取 title 转文本
                t = re.search(r"<title[^>]*>(.*?)</title>", token, re.S)
                latex = _latex_to_text(t.group(1)) if t else ""
                if latex:
                    p.add_run(" " + latex + " ")
            pos = m.end()
        # 剩余文本
        tail = re.sub(r"<[^>]+>", "", seg[pos:])
        tail = html.unescape(tail)
        if tail.strip():
            p.add_run(tail)
        paras.append(p)
    return paras


def export_docx(questions, out_path, title="", img_dir=".import_imgs", instructions=None):
    """题目列表 → docx。结构：标题/格式说明 + 1. 题干 / 【答案】 / 【解析】，图片回嵌。"""
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    if instructions:
        for ins in instructions:
            p = doc.add_paragraph()
            p.add_run(ins)
    for i, q in enumerate(questions, 1):
        # 题号 + 题干
        p = doc.add_paragraph()
        run = p.add_run("%d. " % i)
        run.bold = True
        p.add_run("【%s】%s" % (q.get("type") or "解答题", ""))
        _html_to_paras(doc, q.get("stem", ""))
        # 选项
        for o in (q.get("options") or []):
            doc.add_paragraph(_coerce_str(o))
        # 答案（含图片回嵌）
        if q.get("answer"):
            ap = doc.add_paragraph()
            ar = ap.add_run("【答案】")
            ar.bold = True
            _html_to_paras(doc, q.get("answer"))
        # 解析（含图片回嵌）
        if q.get("analysis"):
            xp = doc.add_paragraph()
            xr = xp.add_run("【解析】")
            xr.bold = True
            _html_to_paras(doc, q.get("analysis"))
    doc.save(out_path)
    return out_path


def _strip_html(s):
    s = re.sub(r"<[^>]+>", "", s or "")
    return html.unescape(s).strip()


# ---------------------------------------------------------------- 表格模板（人手友好）
_TEMPLATE_HEADERS = ["题型", "题干", "选项", "答案", "解析", "难度(1-5)", "学科", "年级", "知识点"]
_TYPE_CN = {"单选": "single_choice", "多选": "multi_choice", "填空": "fill_blank",
            "判断": "true_false", "解答": "essay", "解答题": "essay", "简答": "essay"}


def _shade_cell(cell, color):
    """给表格单元格加底色（浅黄=示例行，浅蓝=表头）。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def export_table_template_docx(out_path, title="题库导入模板", empty_rows=15):
    """生成"人手可操作"的表格模板：填写说明 + 表头 + 示例行 + 空行。
    导入解析(_parse_table_template)可直接识别本模板。"""
    from docx.shared import Pt, RGBColor
    doc = Document()
    doc.add_heading(title, level=1)

    ins = [
        "【使用方法】本表每行一道题，按表头填写即可；填好后保存，在「题库管理 → 导入管理」上传本文件，系统自动识别入库。",
        "【题型】只填：单选 / 多选 / 填空 / 判断 / 解答（其他文字无效）。",
        "【题干】题目完整文字；数学公式用括号写，如 a/(b+c)、√2、x^2。",
        "【选项】只有单选/多选需要填；选项之间用 | 分隔，例：A.甲 | B.乙 | C.丙 | D.丁。",
        "【答案】单选/多选填选项字母（如 A 或 AC）；判断填 正确/错误；填空填答案；解答填参考答案要点。",
        "【难度】1~5 数字（1 最简单、5 最难），不填默认 3。",
        "【学科/年级】如 数学、语文、初一、初二；不填则按导入页选择。",
        "【知识点】选填，多个用 | 分隔。",
        "【示例】下面黄色行是填写示例，按它的格式填即可；示例行可保留（会一并导入）也可删除。",
    ]
    for s in ins:
        p = doc.add_paragraph()
        r = p.add_run(s)
        if s.startswith("【"):
            r.bold = True

    # 示例行（黄色）
    sample = [
        ["单选", "1+1 等于几？", "A.1 | B.2 | C.3 | D.4", "B", "因为 1+1=2", "3", "数学", "初一", "有理数"],
        ["填空", "5 的相反数是____", "", "-5", "相反数概念", "2", "数学", "初一", "相反数"],
        ["判断", "0 是最小的正整数。", "", "错误", "0 不是正数", "2", "数学", "初一", "有理数"],
    ]
    table = doc.add_table(rows=1 + len(sample) + empty_rows, cols=len(_TEMPLATE_HEADERS))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    # 表头（浅蓝 + 加粗）
    for j, h in enumerate(_TEMPLATE_HEADERS):
        cell = table.cell(0, j)
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
        _shade_cell(cell, "DEEBF7")
    # 示例行（浅黄）
    for i, row in enumerate(sample, 1):
        for j, v in enumerate(row):
            cell = table.cell(i, j)
            cell.text = v or ""
            _shade_cell(cell, "FFF2CC")
    # 其余空行留白（不填任何内容）
    doc.save(out_path)
    return out_path


def _parse_table_template(blocks, default_subject="", default_grade="", default_difficulty=3,
                          source="docx", category="", src_prefix=""):
    """识别表格模板（表头含 题型+题干）→ 逐行解析为系统题目 JSON；不是模板返回 None。"""
    for kind, payload in blocks:
        if kind != "table" or not payload or not payload[0]:
            continue
        # 新格式 payload=(html, imgs)；模板识别需要纯文本行。从 HTML 反解或用旧式 rows。
        # 用正则从 HTML 提取每个 <td> 文本作为"行"（首个 <tr> 为表头）
        html_txt = payload[0]
        trs = re.findall(r"<tr>(.*?)</tr>", html_txt, re.S)
        if not trs:
            continue
        head = [re.sub(r"<[^>]+>", "", c).strip() for c in re.findall(r"<td>(.*?)</td>", trs[0], re.S)]
        if not any("题型" in h for h in head) or not any("题干" in h for h in head):
            continue
        idx = {}
        for j, h in enumerate(head):
            if "题型" in h: idx["type"] = j
            elif "题干" in h: idx["stem"] = j
            elif "选项" in h: idx["options"] = j
            elif "答案" in h: idx["answer"] = j
            elif "解析" in h: idx["analysis"] = j
            elif "难度" in h: idx["difficulty"] = j
            elif "学科" in h: idx["subject"] = j
            elif "年级" in h: idx["grade"] = j
            elif "知识点" in h: idx["knowledge"] = j
        if "type" not in idx or "stem" not in idx:
            continue

        def _cell(r, key):
            j = idx.get(key)
            if j is None or j >= len(r):
                return ""
            return str(r[j]).strip()

        out = []
        # 数据行：trs[1:]（trs[0] 是表头）；单元格反解时保留 <img> 标签（图片内容不能丢）
        _cell_text = lambda c: re.sub(r"<[^>]+>", "", c).strip()
        for num, tr in enumerate(trs[1:], 1):
            r = []
            for c in re.findall(r"<td>(.*?)</td>", tr, re.S):
                # 保留 <img src=...>，去掉其它标签；多个 img 用换行分隔
                imgs = re.findall(r"<img[^>]*>", c)
                txt = _cell_text(c)
                r.append((txt + ("\n" + "\n".join(imgs) if imgs else "")).strip())
            stem = _cell(r, "stem")
            if not stem:
                continue
            typ_raw = _cell(r, "type")
            typ = _TYPE_CN.get(typ_raw, detect_type(stem, []))
            # 选项：按 | ; ； ， 分隔
            opts_raw = _cell(r, "options")
            opts = [o.strip() for o in re.split(r"[|；;，,]", opts_raw) if o.strip()]
            answer = _cell(r, "answer")
            analysis = _cell(r, "analysis")
            diff = _cell(r, "difficulty")
            try:
                diff = int(diff) if diff else default_difficulty
            except ValueError:
                diff = default_difficulty
            subj = _cell(r, "subject") or default_subject
            grd = _cell(r, "grade") or default_grade
            kp = [k.strip() for k in re.split(r"[|；;，,]", _cell(r, "knowledge")) if k.strip()]
            # 通用：文本转义、<img> 标签原样保留（题干/答案/解析都可能含图）
            def _html_keep_img(text):
                return "".join(
                    html.escape(part) if i % 2 == 0 else part
                    for i, part in enumerate(re.split(r"(<img[^>]*>)", text))
                ).replace("\n", "</p><p>")
            out.append({
                "id": "IMP%03d" % num,
                "code": _gen_code(subj, grd, num),
                "subject": subj,
                "grade": grd,
                "type": typ,
                "difficulty": diff,
                "score": 0,
                "stem": "<p>" + _html_keep_img(stem) + "</p>",
                "options": opts,
                "answer": "<p>" + _html_keep_img(answer) + "</p>",
                "analysis": "<p>" + _html_keep_img(analysis) + "</p>",
                "knowledge": kp,
                "category": category,
                "source": source,
                "images": [],
            })
        if out:
            return out
    return None


def _coerce_str(v):
    """把 answer/analysis/stem 等字段归一化为字符串（兼容 list/dict/None）。"""
    if v is None:
        return ""
    if isinstance(v, str):
        return v
    if isinstance(v, (list, tuple)):
        return "\n".join(_coerce_str(x) for x in v if x not in (None, ""))
    if isinstance(v, dict):
        return json.dumps(v, ensure_ascii=False)
    return str(v)


# ---------------------------------------------------------------- CLI
if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser(description="题库 docx 导入/导出")
    ap.add_argument("--check", help="格式检查一个 docx")
    ap.add_argument("--import", dest="imp", help="导入 docx，输出 JSON")
    ap.add_argument("--export", help="从 JSON 导出 docx")
    ap.add_argument("--img", default=".import_imgs", help="图片目录")
    ap.add_argument("--subject", default="")
    ap.add_argument("--grade", default="")
    ap.add_argument("--out", default="out")
    a = ap.parse_args()
    if a.check:
        print(json.dumps(check_docx(a.check, a.img), ensure_ascii=False, indent=2))
    elif a.imp:
        qs = parse_docx(a.imp, a.img, a.subject, a.grade)
        with open(a.out if a.out.endswith(".json") else a.out + ".json", "w", encoding="utf-8") as f:
            json.dump(qs, f, ensure_ascii=False, indent=2)
        print("导入 %d 题 -> %s" % (len(qs), a.out))
    elif a.export:
        with open(a.export, encoding="utf-8") as f:
            qs = json.load(f)
        export_docx(qs, a.out if a.out.endswith(".docx") else a.out + ".docx")
        print("导出 -> %s" % a.out)
